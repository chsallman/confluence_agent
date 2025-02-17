import csv
import io
import os

import docx  # python-docx for DOCX files
import fitz  # PyMuPDF for PDFs
import pandas as pd
import requests
from atlassian import Confluence
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv

from logger_config import setup_logger

# Initialize logger
logger = setup_logger("my_logger")

# Example logs

load_dotenv()


def take_data_from_confluence():
    try:
        logger.info("Connecting to confluence")
        API_TOKEN = os.environ.get("API_TOKEN")
        USERNAME = os.environ.get("USERNAME")
        BASE_URL = os.environ.get("BASE_URL")
        SPACE_KEY = os.environ.get("SPACE_KEY")

        confluence = Confluence(
            url=BASE_URL, username=USERNAME, password=API_TOKEN, cloud=True
        )

        pages = confluence.get_all_pages_from_space(SPACE_KEY, start=0, limit=50)
    except Exception as e:
        logger.error(f"Error occured while taking data from Confluence: {e}")

    return confluence, pages


def take_out_table(soup, page):
    logger.info("Processing Tables")

    all_tables_page = []
    tables = soup.find_all("table")

    for table in tables:
        table_data = []
        rows = table.find_all("tr")

        for row in rows:
            cols = row.find_all(["th", "td"])
            cols = [col.get_text(strip=True) for col in cols]
            table_data.append(cols)
        heading = table.find_previous(["h1", "h2", "h3", "h4", "h5", "h6"])
        heading_text = heading.get_text(strip=True) if heading else None
        if heading_text:
            table_data.insert(0, heading_text)
        all_tables_page.append(
            {"title": page["title"], "id": page["id"], "text": table_data}
        )

    return all_tables_page


def takes_out_plain_text(soup, page):
    logger.info("Processing Text")

    for table in soup.find_all("table"):
        table.decompose()  # Remove tables

    plain_text = soup.get_text(separator="\n", strip=True)

    return {"title": page["title"], "id": page["id"], "text": plain_text}


def extract_attachment_text(file_name, file_url):
    logger.info("Processing Attachments")
    BASE_URL = "https://techcorpdocumentation.atlassian.net/wiki"
    download_url = BASE_URL + file_url
    API_TOKEN = os.environ.get("API_TOKEN")
    USERNAME = os.environ.get("USERNAME")

    response = requests.get(download_url, auth=(USERNAME, API_TOKEN))

    if response.status_code != 200:
        return f"Failed to download. Status code: {response.status_code}"

    # Determine file type
    file_extension = file_name.lower().split(".")[-1]

    # Process PDF
    try:
        if file_extension == "pdf":
            temp_file = "temp_attachment.pdf"
            with open(temp_file, "wb") as file:
                file.write(response.content)

            text = []
            with fitz.open(temp_file) as doc:
                for page in doc:
                    text.append(page.get_text())

            os.remove(temp_file)  # Cleanup
            return "\n".join(text) if text else "No readable text found."

        # Process DOCX
        elif file_extension == "docx":
            temp_file = "temp_attachment.docx"
            with open(temp_file, "wb") as file:
                file.write(response.content)

            doc = docx.Document(temp_file)
            os.remove(temp_file)  # Cleanup
            return (
                "\n".join([para.text for para in doc.paragraphs])
                if doc.paragraphs
                else "No readable text found."
            )

        # Process CSV
        elif file_extension == "csv":
            text = []
            csv_data = io.StringIO(response.content.decode("utf-8"))
            reader = csv.reader(csv_data)
            for row in reader:
                text.append(", ".join(row))

            return "\n".join(text) if text else "No readable text found."
    except Exception as e:
        logger.error(f"Error processing {file_name}: {e}")
    return "Unsupported file type."


def get_attachments(confluence, page):
    logger.info("Getting Attachments")

    """Retrieve and extract attachments from pages."""
    page_id = page["id"]
    attachments = confluence.get_attachments_from_content(page_id)
    extracted_attachments = []

    for attachment in attachments["results"]:
        file_name = attachment["title"]
        file_url = attachment["_links"]["download"]
        # base_url = "https://techcorpdocumentation.atlassian.net/wiki"
        text_data = extract_attachment_text(file_name, file_url)
        if text_data:
            final_text = {
                "title": page["title"],
                "id": page["id"],
                "text": text_data,
                "file_name": file_name,
                "file_url": file_url,
            }
            extracted_attachments.append(final_text)

    return extracted_attachments


def take_data_from_pages(confluence, pages):
    logger.info("Extracting Data from pages")

    all_tables = []
    all_text_data = []
    all_attachments = []

    for page in pages:
        page_id = page["id"]
        page_content = confluence.get_page_by_id(page_id, expand="body.storage")
        html_content = page_content["body"]["storage"]["value"]
        soup = BeautifulSoup(html_content, "html.parser")

        all_tables.extend(take_out_table(soup, page))
        all_text_data.append(takes_out_plain_text(soup, page))
        all_attachments.extend(get_attachments(confluence, page))

    return all_tables, all_text_data, all_attachments
